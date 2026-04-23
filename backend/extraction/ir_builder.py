"""IR Builder: construye Intermediate Representation desde documentos.

Usa fitz (PyMuPDF) para PDFs y python-docx para DOCX.
Detecta zonas semanticas: HEADER, RADICADO, PARTIES, DATES, BODY, FOOTER, etc.
"""

import logging
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from sqlalchemy.orm import Session

from backend.database.models import Case, Email
from backend.extraction.ir_models import (
    TextSpan, DocumentZone, DocumentIR, CaseIR,
)
from backend.extraction.pipeline import classify_doc_type
from backend.extraction.pdf_visual_analyzer import analyze_pdf_visual, report_to_zone_metadata

logger = logging.getLogger("tutelas.ir_builder")


def _libreoffice_convert_to_docx(src: Path) -> Path | None:
    """Convierte .doc/.rtf/.odt → .docx con LibreOffice headless. None si falla."""
    if not shutil.which("libreoffice"):
        return None
    try:
        tmpdir = Path(tempfile.mkdtemp(prefix="lo_convert_"))
        r = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx",
             "--outdir", str(tmpdir), str(src)],
            capture_output=True, timeout=30,
        )
        if r.returncode != 0:
            return None
        converted = tmpdir / (src.stem + ".docx")
        return converted if converted.exists() else None
    except Exception as e:
        logger.debug("LibreOffice convert falló para %s: %s", src.name, e)
        return None


def _extract_text_antiword(src: Path) -> str:
    """Texto plano de .doc binario via antiword (paquete apt). '' si no disponible."""
    if not shutil.which("antiword"):
        return ""
    try:
        r = subprocess.run(["antiword", str(src)], capture_output=True, timeout=20)
        return r.stdout.decode("utf-8", errors="replace")
    except Exception:
        return ""


def _extract_text_olefile_scan(src: Path) -> str:
    """Último recurso para .doc binario corrupto: escanea streams OLE por ASCII."""
    try:
        import olefile
        with olefile.OleFileIO(str(src)) as ole:
            for stream in ("WordDocument", "1Table", "0Table"):
                if ole.exists(stream):
                    raw = ole.openstream(stream).read()
                    chunks = re.findall(rb"[\x20-\x7e\xc0-\xff\s]{10,}", raw)
                    decoded = b" ".join(chunks[:500]).decode("utf-8", errors="replace")
                    if len(decoded.strip()) > 100:
                        return decoded
    except Exception:
        pass
    return ""

_IR_BODY_MAX = 150000   # Guard-rail amplio para zona BODY (full_text siempre se preserva)
_FOOTER_TAIL_CHARS = 4000  # Últimos N chars preservados íntegros como zona FOOTER_TAIL


def _make_body_zone(text: str, filename: str = "", **kwargs) -> DocumentZone:
    """Crear zona BODY. El truncamiento aquí es solo guard-rail de RAM; los
    extractores locales consumen DocumentIR.full_text (no truncado)."""
    truncated = len(text) > _IR_BODY_MAX
    if truncated:
        logger.debug("BODY zone trimmed de %d a %d chars para %s (full_text intacto)",
                     len(text), _IR_BODY_MAX, filename)
    return DocumentZone(
        zone_type="BODY", text=text[:_IR_BODY_MAX],
        truncated=truncated, **kwargs,
    )


def _make_footer_tail_zone(full_text: str, page: int = 0) -> DocumentZone | None:
    """Zona FOOTER_TAIL: últimos 4K chars del documento, íntegros.

    Garantiza que 'Proyectó: X', sellos finales y firmas de abogado queden
    siempre disponibles para regex/forensic aunque el doc sea gigante.
    """
    if not full_text or len(full_text.strip()) < 100:
        return None
    tail = full_text[-_FOOTER_TAIL_CHARS:] if len(full_text) > _FOOTER_TAIL_CHARS else full_text
    return DocumentZone(
        zone_type="FOOTER_TAIL", text=tail, page=page, confidence=1.0,
        metadata={"total_doc_chars": len(full_text)},
    )


# ---------------------------------------------------------------------------
# Patrones para deteccion de zonas
# ---------------------------------------------------------------------------

_RE_RADICADO_23 = re.compile(r"(68[\d]{17,21})")
_RE_RADICADO_SEP = re.compile(
    r"(68[\d]{3,5}[-\s\.]?\d{2}[-\s\.]?\d{2}[-\s\.]?\d{3}[-\s\.]?\d{4}[-\s\.]?\d{5}[-\s\.]?\d{2})"
)
_RE_FECHA_ESCRITA = re.compile(
    r"(\d{1,2})\s*(?:\(\w+\))?\s*(?:d[eí]as?\s+)?(?:del?\s+mes\s+de\s+)?"
    r"(enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre)"
    r"\s+(?:del?\s+)?(?:a[ñn]o\s+)?(?:dos\s+mil\s+veinti\w+|\(?(20\d{2})\)?)",
    re.IGNORECASE,
)
_RE_FECHA_NUM = re.compile(r"(\d{1,2})[/\-](\d{1,2})[/\-](20\d{2})")
_RE_PARTIES_START = re.compile(
    r"(?i)(accionante|demandante|promovid[ao]|parte\s+actora)\s*[:.\-]?"
)
_RE_RIGHTS = re.compile(
    r"(?i)(derecho[s]?\s+(?:fundamental|vulnerad|invocad))"
)
_RE_RESUELVE = re.compile(r"(?i)^\s*(RESUELVE|FALLA|DECIDE)\s*[:.]?\s*$")
_RE_FOREST = re.compile(r"(?:FOREST|forest|Forest)\s*(?:No\.?\s*)?:?\s*(\d{5,13})")
_RE_FOREST_PHRASE = re.compile(
    r"(?:n[uú]mero\s+de\s+radicado\s+es|radicado\s+(?:y\s+enviado|es))\s+(\d{7,13})",
    re.IGNORECASE,
)

_MESES = {
    "enero": "01", "febrero": "02", "marzo": "03", "abril": "04",
    "mayo": "05", "junio": "06", "julio": "07", "agosto": "08",
    "septiembre": "09", "octubre": "10", "noviembre": "11", "diciembre": "12",
}

# Prioridades de doc_type para ordenamiento
_DOC_PRIORITY = {
    "PDF_AUTO_ADMISORIO": 1, "PDF_SENTENCIA": 2, "DOCX_RESPUESTA": 3,
    "PDF_GMAIL": 4, "EMAIL_MD": 4, "EMAIL_DB": 4,
    "PDF_IMPUGNACION": 5, "PDF_INCIDENTE": 6,
    "DOCX_DESACATO": 6, "DOCX_IMPUGNACION": 5, "DOCX_CUMPLIMIENTO": 7,
}


# ---------------------------------------------------------------------------
# PDF IR Builder (usando fitz)
# ---------------------------------------------------------------------------

def _build_pdf_ir(file_path: str, doc_type: str) -> DocumentIR:
    """Construir IR desde un PDF usando fitz (PyMuPDF)."""
    import fitz

    path = Path(file_path)
    if not path.exists():
        return DocumentIR(
            filename=path.name, doc_type=doc_type,
            priority=_DOC_PRIORITY.get(doc_type, 9),
        )

    try:
        doc = fitz.open(str(path))
    except Exception as e:
        logger.warning("fitz no pudo abrir %s: %s", path.name, e)
        return DocumentIR(
            filename=path.name, doc_type=doc_type,
            priority=_DOC_PRIORITY.get(doc_type, 9),
        )

    zones = []
    full_text_parts = []
    page_count = len(doc)
    has_ocr = False

    for page_num in range(page_count):
        page = doc[page_num]
        page_height = page.rect.height
        page_dict = page.get_text("dict")
        page_text = page.get_text("text")
        full_text_parts.append(page_text)

        if len(page_text.strip()) < 50:
            has_ocr = True

        # Recolectar spans con metadata
        page_spans = []
        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:  # solo bloques de texto
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if not text:
                        continue
                    ts = TextSpan(
                        text=text,
                        font_name=span.get("font", ""),
                        font_size=span.get("size", 0),
                        is_bold="Bold" in span.get("font", "") or "bold" in span.get("font", "") or bool(span.get("flags", 0) & 16),
                        is_italic="Italic" in span.get("font", "") or "italic" in span.get("font", ""),
                        page=page_num + 1,
                    )
                    y_pos = span.get("origin", (0, 0))[1]
                    page_spans.append((ts, y_pos, block.get("bbox", [0, 0, 0, 0])))

        # --- Deteccion de zonas por posicion y contenido ---

        # HEADER: fuente grande o bold en el top 25% de pagina 1
        if page_num == 0:
            header_spans = []
            header_texts = []
            for ts, y_pos, bbox in page_spans:
                in_top = y_pos < page_height * 0.35
                is_prominent = ts.font_size >= 10 and (ts.is_bold or ts.font_size >= 12)
                if in_top and is_prominent:
                    header_spans.append(ts)
                    header_texts.append(ts.text)

            if header_texts:
                header_text = " ".join(header_texts)
                zones.append(DocumentZone(
                    zone_type="HEADER", text=header_text,
                    spans=header_spans, page=1, confidence=0.9,
                ))

        # RADICADO: buscar patron de 23 digitos en toda la pagina
        for ts, y_pos, bbox in page_spans:
            m = _RE_RADICADO_23.search(ts.text) or _RE_RADICADO_SEP.search(ts.text)
            if m:
                zones.append(DocumentZone(
                    zone_type="RADICADO", text=m.group(1),
                    spans=[ts], page=page_num + 1, confidence=0.95,
                    metadata={"raw_match": m.group(0)},
                ))
                break  # solo el primero por pagina

        # DATES: buscar fechas escritas o numericas
        for ts, y_pos, bbox in page_spans:
            m_esc = _RE_FECHA_ESCRITA.search(ts.text)
            if m_esc:
                dia = m_esc.group(1).zfill(2)
                mes = _MESES.get(m_esc.group(2).lower(), "00")
                anio = m_esc.group(3) if m_esc.group(3) else "2026"
                fecha = f"{dia}/{mes}/{anio}"
                zones.append(DocumentZone(
                    zone_type="DATES", text=ts.text,
                    page=page_num + 1, confidence=0.85,
                    metadata={"fecha_parsed": fecha, "format": "escrita"},
                ))
                break
            m_num = _RE_FECHA_NUM.search(ts.text)
            if m_num:
                fecha = f"{m_num.group(1).zfill(2)}/{m_num.group(2).zfill(2)}/{m_num.group(3)}"
                zones.append(DocumentZone(
                    zone_type="DATES", text=ts.text,
                    page=page_num + 1, confidence=0.90,
                    metadata={"fecha_parsed": fecha, "format": "numerica"},
                ))
                break

        # PARTIES: buscar seccion accionante/demandante (pagina 1-2)
        if page_num < 2:
            for ts, y_pos, bbox in page_spans:
                if _RE_PARTIES_START.search(ts.text):
                    # Recolectar texto desde aqui hasta la siguiente seccion
                    parties_text = []
                    capturing = False
                    for ts2, _, _ in page_spans:
                        if _RE_PARTIES_START.search(ts2.text):
                            capturing = True
                        if capturing:
                            parties_text.append(ts2.text)
                            if len(parties_text) > 20:
                                break
                    if parties_text:
                        zones.append(DocumentZone(
                            zone_type="PARTIES", text=" ".join(parties_text),
                            page=page_num + 1, confidence=0.80,
                        ))
                    break

        # RIGHTS: buscar derechos vulnerados
        if page_num < 3:
            for ts, y_pos, bbox in page_spans:
                if _RE_RIGHTS.search(ts.text):
                    rights_text = [ts.text]
                    zones.append(DocumentZone(
                        zone_type="RIGHTS", text=" ".join(rights_text),
                        page=page_num + 1, confidence=0.75,
                    ))
                    break

        # RESOLUTION: buscar RESUELVE/FALLA (sentencias)
        if doc_type in ("PDF_SENTENCIA", "PDF_INCIDENTE"):
            for ts, y_pos, bbox in page_spans:
                if _RE_RESUELVE.search(ts.text):
                    resolution_texts = []
                    capture = False
                    for ts2, _, _ in page_spans:
                        if _RE_RESUELVE.search(ts2.text):
                            capture = True
                        if capture:
                            resolution_texts.append(ts2.text)
                            if len(resolution_texts) > 30:
                                break
                    if resolution_texts:
                        zones.append(DocumentZone(
                            zone_type="RESOLUTION",
                            text=" ".join(resolution_texts),
                            page=page_num + 1, confidence=0.85,
                        ))
                    break

        # FOOTER: ultimo 15% de cada pagina
        footer_spans = []
        for ts, y_pos, bbox in page_spans:
            if y_pos > page_height * 0.85:
                footer_spans.append(ts)
        if footer_spans:
            footer_text = " ".join(s.text for s in footer_spans)
            if len(footer_text.strip()) > 5:
                zones.append(DocumentZone(
                    zone_type="FOOTER", text=footer_text,
                    spans=footer_spans, page=page_num + 1,
                ))

    # Extraer tablas (max 5 tablas utiles por documento)
    tables = []
    try:
        for page_num in range(min(page_count, 10)):
            if len(tables) >= 5:
                break
            page = doc[page_num]
            page_tables = page.find_tables()
            for t in page_tables.tables:
                rows = t.extract()
                if not rows or len(rows) < 2:
                    continue
                # Validar que la tabla tiene contenido real
                all_text = " ".join(str(c or "") for r in rows for c in r)
                if len(all_text.strip()) < 20:
                    continue  # Tabla vacia o solo formato
                tables.append(rows)
    except Exception:
        pass

    # BODY: texto de paginas principales (sin duplicar header/footer)
    body_text = "\n".join(full_text_parts)
    if body_text.strip():
        zones.append(_make_body_zone(body_text, filename=path.name, page=0, confidence=1.0))
        tail_zone = _make_footer_tail_zone(body_text, page=page_count)
        if tail_zone:
            zones.append(tail_zone)

    doc.close()

    # v5.5: análisis visual determinista (sin IA) — detecta logos, sellos, watermarks,
    # firmas y texto rotado. Enriquece el IR con señales de institucionalidad.
    try:
        visual = analyze_pdf_visual(str(path))
        if visual.findings or visual.images_count > 0:
            rotated_text = "\n".join(visual.rotated_text_snippets[:10])
            zones.append(DocumentZone(
                zone_type="VISUAL",
                text=rotated_text or f"[{visual.images_count} imgs, {visual.annotations_count} annots]",
                page=0,
                confidence=visual.institutional_score,
                metadata=report_to_zone_metadata(visual),
            ))
    except Exception as e:
        logger.debug("Visual analyzer falló para %s: %s", path.name, e)

    return DocumentIR(
        filename=path.name,
        doc_type=doc_type,
        priority=_DOC_PRIORITY.get(doc_type, 9),
        zones=zones,
        tables=tables,
        full_text=body_text,
        page_count=page_count,
        has_ocr_pages=has_ocr,
        extraction_method="fitz",
    )


# ---------------------------------------------------------------------------
# DOCX IR Builder
# ---------------------------------------------------------------------------

def _build_docx_ir(file_path: str, doc_type: str) -> DocumentIR:
    """Construir IR desde un DOCX usando python-docx."""
    import docx
    path = Path(file_path)
    if not path.exists():
        return DocumentIR(
            filename=path.name, doc_type=doc_type,
            priority=_DOC_PRIORITY.get(doc_type, 9),
        )

    try:
        doc = docx.Document(str(path))
    except Exception as e:
        logger.warning("python-docx no pudo abrir %s: %s — intentando fallbacks", path.name, e)
        # Cadena de fallback: LibreOffice convert → antiword → olefile scan.
        converted = _libreoffice_convert_to_docx(path)
        if converted:
            try:
                doc = docx.Document(str(converted))
                logger.info("Recuperado vía LibreOffice: %s", path.name)
                return _build_docx_from_open(doc, path, doc_type, method="libreoffice+docx")
            except Exception as e2:
                logger.debug("LibreOffice→python-docx falló para %s: %s", path.name, e2)
        # Texto plano sin estructura DOCX
        plain = _extract_text_antiword(path) or _extract_text_olefile_scan(path)
        if plain and len(plain.strip()) > 100:
            logger.info("Recuperado vía texto plano (%d chars): %s", len(plain), path.name)
            zones = [_make_body_zone(plain, filename=path.name, confidence=0.75)]
            tail = _make_footer_tail_zone(plain)
            if tail:
                zones.append(tail)
            return DocumentIR(
                filename=path.name, doc_type=doc_type,
                priority=_DOC_PRIORITY.get(doc_type, 9),
                zones=zones, full_text=plain, page_count=1,
                extraction_method="fallback_plain",
            )
        return DocumentIR(
            filename=path.name, doc_type=doc_type,
            priority=_DOC_PRIORITY.get(doc_type, 9),
        )

    return _build_docx_from_open(doc, path, doc_type, method="python-docx")


def _build_docx_from_open(doc, path: Path, doc_type: str, method: str = "python-docx") -> DocumentIR:
    """Construye el IR a partir de un objeto python-docx ya abierto.

    Se extrae para reutilizarse desde el flujo directo y desde el fallback
    LibreOffice→python-docx.
    """
    zones = []
    full_text_parts = []

    # Headers
    for section in doc.sections:
        if section.header and section.header.paragraphs:
            header_text = " ".join(p.text for p in section.header.paragraphs if p.text.strip())
            if header_text.strip():
                zones.append(DocumentZone(
                    zone_type="HEADER", text=header_text, confidence=0.9,
                ))
                # Buscar FOREST en header
                m = _RE_FOREST.search(header_text) or _RE_FOREST_PHRASE.search(header_text)
                if m:
                    zones.append(DocumentZone(
                        zone_type="WATERMARK", text=m.group(1),
                        confidence=0.90,
                        metadata={"forest_number": m.group(1)},
                    ))

    # Body paragraphs con estilos
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        full_text_parts.append(text)

        # Detectar si es heading
        style_name = para.style.name if para.style else ""
        is_heading = "Heading" in style_name or "heading" in style_name

        # Detectar runs bold
        has_bold = any(run.font.bold for run in para.runs if run.font.bold)
        font_size = 0
        for run in para.runs:
            if run.font.size:
                font_size = max(font_size, run.font.size.pt if hasattr(run.font.size, 'pt') else 0)

        span = TextSpan(
            text=text, font_name=style_name,
            font_size=font_size, is_bold=has_bold or is_heading,
        )

        # Clasificar parrafos en zonas
        if _RE_RESUELVE.search(text):
            zones.append(DocumentZone(
                zone_type="RESOLUTION", text=text, spans=[span], confidence=0.85,
            ))
        elif _RE_PARTIES_START.search(text):
            zones.append(DocumentZone(
                zone_type="PARTIES", text=text, spans=[span], confidence=0.80,
            ))

    # Footers — aqui esta el abogado (regular + first_page + even_page)
    footer_texts = set()
    for section in doc.sections:
        for footer_attr in ("footer", "first_page_footer", "even_page_footer"):
            footer = getattr(section, footer_attr, None)
            if footer and hasattr(footer, "paragraphs"):
                ft = " ".join(p.text for p in footer.paragraphs if p.text.strip())
                if ft.strip() and ft not in footer_texts:
                    footer_texts.add(ft)
                    zones.append(DocumentZone(
                        zone_type="FOOTER", text=ft, confidence=0.95,
                    ))

    # Tablas
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append(rows)

    body_text = "\n".join(full_text_parts)
    zones.append(_make_body_zone(body_text, filename=path.name, confidence=1.0))
    tail_zone = _make_footer_tail_zone(body_text)
    if tail_zone:
        zones.append(tail_zone)

    return DocumentIR(
        filename=path.name,
        doc_type=doc_type,
        priority=_DOC_PRIORITY.get(doc_type, 9),
        zones=zones,
        tables=tables,
        full_text=body_text,
        page_count=1,
        extraction_method=method,
    )


# ---------------------------------------------------------------------------
# Email IR Builder
# ---------------------------------------------------------------------------

def _build_email_dict(email: Email) -> dict:
    """Construir dict de email para el IR."""
    return {
        "subject": email.subject or "",
        "sender": email.sender or "",
        "date": email.date_received.strftime("%d/%m/%Y") if email.date_received else "",
        "body": email.body_preview or "",
    }


# ---------------------------------------------------------------------------
# Case IR Builder (punto de entrada principal)
# ---------------------------------------------------------------------------

def build_case_ir(db: Session, case: Case) -> CaseIR:
    """Construir IR completo de un caso desde sus documentos y emails.

    Este es el punto de entrada principal del IR Builder.
    """
    logger.info("Construyendo IR para caso %d: %s", case.id, case.folder_name)

    # Construir IR de cada documento
    doc_irs = []
    for doc in case.documents:
        if not doc.file_path:
            continue
        ext = Path(doc.file_path).suffix.lower()
        doc_type = classify_doc_type(doc.filename)

        if ext == ".pdf":
            ir = _build_pdf_ir(doc.file_path, doc_type)
            # Fallback: si fitz no saca texto (PDF escaneado/OCR-only) pero la DB
            # tiene extracted_text (vino de OCR previo), usar ese texto.
            if (not ir.full_text or len(ir.full_text.strip()) < 50) and doc.extracted_text:
                ir = DocumentIR(
                    filename=doc.filename, doc_type=doc_type,
                    priority=_DOC_PRIORITY.get(doc_type, 9),
                    full_text=doc.extracted_text,
                    zones=[_make_body_zone(doc.extracted_text, filename=doc.filename)],
                    extraction_method="db_fallback_ocr",
                )
        elif ext in (".docx", ".doc"):
            ir = _build_docx_ir(doc.file_path, doc_type)
        elif ext == ".md":
            # Markdown: leer como texto plano
            try:
                text = Path(doc.file_path).read_text(encoding="utf-8", errors="replace")
                ir = DocumentIR(
                    filename=doc.filename, doc_type=doc_type,
                    priority=_DOC_PRIORITY.get(doc_type, 9),
                    full_text=text,
                    zones=[_make_body_zone(text, filename=doc.filename)],
                    extraction_method="markdown",
                )
            except Exception:
                continue
        else:
            # Tipo no soportado: incluir con texto existente si lo tiene
            if doc.extracted_text:
                ir = DocumentIR(
                    filename=doc.filename, doc_type=doc_type,
                    priority=_DOC_PRIORITY.get(doc_type, 9),
                    full_text=doc.extracted_text,
                    zones=[_make_body_zone(doc.extracted_text, filename=doc.filename)],
                    extraction_method="existing",
                )
            else:
                continue

        doc_irs.append(ir)

    # Ordenar por prioridad
    doc_irs.sort(key=lambda d: d.priority)

    # Emails del caso
    emails = db.query(Email).filter(Email.case_id == case.id).all()
    email_dicts = [_build_email_dict(e) for e in emails]

    # Campos ya conocidos en DB
    from backend.database.models import Case as CaseModel
    known = {}
    field_map = getattr(CaseModel, "CSV_FIELD_MAP", {})
    for csv_name, attr in field_map.items():
        val = getattr(case, attr, None)
        if val and str(val).strip():
            known[csv_name] = str(val).strip()

    # Correcciones historicas
    corrections = []
    try:
        from backend.agent.memory import get_recent_corrections
        corrs = get_recent_corrections(db, case_id=case.id, limit=10)
        for c in corrs:
            corrections.append({
                "field": c.field_name,
                "ai_value": c.ai_value,
                "corrected_value": c.corrected_value,
                "case_folder": c.case_folder,
            })
    except Exception:
        pass

    ir = CaseIR(
        case_id=case.id,
        folder_name=case.folder_name or "",
        documents=doc_irs,
        emails=email_dicts,
        known_fields=known,
        corrections=corrections,
    )

    logger.info(
        "IR construido: %d docs, %d zonas totales, %d emails",
        len(doc_irs),
        sum(len(d.zones) for d in doc_irs),
        len(email_dicts),
    )

    return ir
