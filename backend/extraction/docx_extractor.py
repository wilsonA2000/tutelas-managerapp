"""Extractor robusto de DOCX con footers, tablas, metadata y ZIP fallback."""

from dataclasses import dataclass
from pathlib import Path
import shutil
import tempfile
import zipfile
import re


@dataclass
class DOCXExtractionResult:
    text: str
    footer_text: str = ""
    header_text: str = ""
    table_text: str = ""
    metadata: dict | None = None
    method: str = "python-docx"
    error: str | None = None
    lawyer_name: str = ""


def _extract_header_from_zip(file_path: Path) -> str:
    """Leer header XML crudo del DOCX (captura textboxes/watermarks con datos FOREST)."""
    try:
        with zipfile.ZipFile(str(file_path), "r") as zf:
            for xml_name in ["word/header1.xml", "word/header2.xml", "word/header3.xml"]:
                if xml_name in zf.namelist():
                    content = zf.read(xml_name).decode("utf-8", errors="replace")
                    clean = re.sub(r"<[^>]+>", " ", content)
                    clean = re.sub(r"\s+", " ", clean).strip()
                    # Solo retornar si tiene datos utiles (FOREST, Proc, etc.)
                    if any(kw in clean for kw in ["Proc", "GOBERNACIÓN", "Radicadora", "Fecha", "Tercero"]):
                        return f"[HEADER FOREST] {clean}"
        return ""
    except Exception:
        return ""


def _extract_lawyer_from_text(text: str) -> str:
    """Buscar nombre del abogado en patrones de firma.
    PRIORIDAD: Proyectó > Elaboró > Revisó > Aprobó
    El que PROYECTA es el abogado responsable, no el que revisa."""
    # Buscar en orden de prioridad (Proyectó primero)
    priority_patterns = [
        r"(?i)proyect[oó]\s*[:\.]\s*(.+?)(?:\n|$)",
        r"(?i)elabor[oó]\s*[:\.]\s*(.+?)(?:\n|$)",
    ]
    fallback_patterns = [
        r"(?i)revis[oó]\s*[:\.]\s*(.+?)(?:\n|$)",
        r"(?i)aprob[oó]\s*[:\.]\s*(.+?)(?:\n|$)",
        r"(?i)profesional\s+(?:universitario|especializado)\s*[-–]\s*(.+?)(?:\n|$)",
    ]

    def _clean_name(name: str) -> str:
        name = re.sub(r"\s*[-–].*$", "", name)
        name = re.sub(r"\s*(CPS|OPS|CONTRATO|CC\.?\s*\d+).*$", "", name, flags=re.IGNORECASE)
        name = name.strip(" .,;:/")
        return name

    # Primero buscar Proyectó/Elaboró (abogado que redactó)
    for pattern in priority_patterns:
        match = re.search(pattern, text)
        if match:
            name = _clean_name(match.group(1).strip())
            if len(name) > 5 and len(name) < 100:
                return name

    # Fallback: Revisó/Aprobó
    for pattern in fallback_patterns:
        match = re.search(pattern, text)
        if match:
            name = _clean_name(match.group(1).strip())
            if len(name) > 5 and len(name) < 100:
                return name

    return ""


def extract_docx(file_path: str | Path) -> DOCXExtractionResult:
    """Extraer texto COMPLETO de un DOCX incluyendo footers y tablas."""
    file_path = Path(file_path)
    if not file_path.exists():
        return DOCXExtractionResult(text="", error=f"Archivo no existe: {file_path}")

    # Intentar metodo principal (python-docx)
    result = _extract_with_python_docx(file_path)
    if result.error:
        # Fallback: metodo ZIP
        result = _extract_with_zip(file_path)

    # Buscar nombre del abogado en todo el texto
    full_text = f"{result.text}\n{result.footer_text}"
    result.lawyer_name = _extract_lawyer_from_text(full_text)

    return result


def _extract_with_python_docx(file_path: Path) -> DOCXExtractionResult:
    """Extraccion con python-docx."""
    try:
        from docx import Document

        doc = Document(str(file_path))

        # Texto del cuerpo
        body_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                body_parts.append(text)

        # Tablas
        table_parts = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_parts.append(" | ".join(cells))

        # Headers (python-docx puede perder textboxes, complementar con ZIP)
        header_parts = []
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    if text:
                        header_parts.append(text)

        # Si python-docx no capturo header, leer XML crudo (contiene datos FOREST)
        if not header_parts:
            forest_header = _extract_header_from_zip(file_path)
            if forest_header:
                header_parts.append(forest_header)

        # Footers (donde esta el nombre del abogado)
        footer_parts = []
        for section in doc.sections:
            if section.footer:
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        footer_parts.append(text)

        # Metadata
        props = doc.core_properties
        metadata = {
            "author": props.author or "",
            "title": props.title or "",
            "created": props.created.isoformat() if props.created else "",
            "modified": props.modified.isoformat() if props.modified else "",
        }

        body_text = "\n".join(body_parts)
        footer_text = "\n".join(footer_parts)
        header_text = "\n".join(header_parts)
        table_text = "\n".join(table_parts)

        # Texto completo consolidado
        full = body_text
        if table_text and table_text not in body_text:
            full += "\n\n[TABLAS]\n" + table_text
        if footer_text:
            full += "\n\n[PIE DE PAGINA]\n" + footer_text
        if header_text:
            full = "[ENCABEZADO]\n" + header_text + "\n\n" + full

        return DOCXExtractionResult(
            text=full,
            footer_text=footer_text,
            header_text=header_text,
            table_text=table_text,
            metadata=metadata,
            method="python-docx",
        )

    except Exception as e:
        return DOCXExtractionResult(text="", error=str(e), method="python-docx")


def _extract_with_zip(file_path: Path) -> DOCXExtractionResult:
    """Fallback: extraer DOCX como ZIP y leer XMLs."""
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            zip_path = Path(tmp_dir) / "doc.zip"
            shutil.copy2(file_path, zip_path)

            parts = {}
            xml_files = ["word/document.xml", "word/footer1.xml", "word/footer2.xml",
                         "word/header1.xml", "word/header2.xml"]

            with zipfile.ZipFile(zip_path, "r") as zf:
                for xml_file in xml_files:
                    if xml_file in zf.namelist():
                        content = zf.read(xml_file).decode("utf-8", errors="replace")
                        # Remover tags XML
                        clean = re.sub(r"<[^>]+>", " ", content)
                        clean = re.sub(r"\s+", " ", clean).strip()
                        parts[xml_file] = clean

            body = parts.get("word/document.xml", "")
            footer = " ".join(
                parts.get(f, "") for f in ["word/footer1.xml", "word/footer2.xml"]
            ).strip()
            header = " ".join(
                parts.get(f, "") for f in ["word/header1.xml", "word/header2.xml"]
            ).strip()

            full = body
            if footer:
                full += "\n\n[PIE DE PAGINA]\n" + footer
            if header:
                full = "[ENCABEZADO]\n" + header + "\n\n" + full

            return DOCXExtractionResult(
                text=full,
                footer_text=footer,
                header_text=header,
                method="zip_fallback",
            )

    except Exception as e:
        return DOCXExtractionResult(text="", error=f"ZIP fallback fallo: {e}", method="zip_fallback")
