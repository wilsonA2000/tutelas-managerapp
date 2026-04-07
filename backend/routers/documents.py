"""Router de documentos."""

from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse, HTMLResponse
from sqlalchemy.orm import Session

from backend.database.database import get_db
from backend.database.models import Document
from backend.services.extraction_service import reextract_doc

router = APIRouter(prefix="/api/documents", tags=["documents"])


@router.get("/{doc_id}")
def api_get_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    data = doc.to_dict()
    data["extracted_text"] = doc.extracted_text or ""
    return data


def _docx_to_html(file_path: str) -> str:
    """Convertir DOCX a HTML para visualizar en navegador."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)

        html_parts = []
        html_parts.append("""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="utf-8">
<style>
  body { font-family: 'Segoe UI', Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 0 20px; color: #333; line-height: 1.6; }
  h1 { font-size: 1.4em; color: #1A5276; border-bottom: 2px solid #1A5276; padding-bottom: 8px; }
  h2 { font-size: 1.2em; color: #2E86C1; }
  table { border-collapse: collapse; width: 100%; margin: 12px 0; }
  td, th { border: 1px solid #ddd; padding: 8px; text-align: left; font-size: 0.9em; }
  th { background: #f5f5f5; font-weight: 600; }
  p { margin: 6px 0; }
  .bold { font-weight: bold; }
  .center { text-align: center; }
  .footer { margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; color: #888; font-size: 0.85em; }
</style>
</head>
<body>""")

        # Procesar parrafos
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name.lower() if para.style else ""
            if "heading 1" in style or "título" in style:
                html_parts.append(f"<h1>{text}</h1>")
            elif "heading 2" in style:
                html_parts.append(f"<h2>{text}</h2>")
            elif para.runs and any(r.bold for r in para.runs):
                html_parts.append(f"<p class='bold'>{text}</p>")
            elif para.alignment and para.alignment == 1:  # CENTER
                html_parts.append(f"<p class='center'>{text}</p>")
            else:
                html_parts.append(f"<p>{text}</p>")

        # Procesar tablas
        for table in doc.tables:
            html_parts.append("<table>")
            for i, row in enumerate(table.rows):
                html_parts.append("<tr>")
                tag = "th" if i == 0 else "td"
                for cell in row.cells:
                    html_parts.append(f"<{tag}>{cell.text.strip()}</{tag}>")
                html_parts.append("</tr>")
            html_parts.append("</table>")

        # Footer
        try:
            for section in doc.sections:
                footer = section.footer
                if footer and footer.paragraphs:
                    footer_text = " | ".join(p.text.strip() for p in footer.paragraphs if p.text.strip())
                    if footer_text:
                        html_parts.append(f"<div class='footer'>{footer_text}</div>")
        except Exception:
            pass

        html_parts.append("</body></html>")
        return "\n".join(html_parts)

    except Exception as e:
        return f"""<!DOCTYPE html><html><body>
        <h2 style='color:#c0392b'>Error al renderizar documento</h2>
        <p>{str(e)}</p>
        </body></html>"""


def _doc_to_html(file_path: str, extracted_text: str = "") -> str:
    """Convertir DOC antiguo a HTML. Intenta antiword/catdoc, si no usa texto extraido de DB."""
    text = ""

    # Intentar herramientas externas
    import shutil
    for tool in ["antiword", "catdoc"]:
        if shutil.which(tool):
            try:
                import subprocess
                result = subprocess.run([tool, file_path], capture_output=True, text=True, timeout=10)
                if result.returncode == 0 and result.stdout.strip():
                    text = result.stdout
                    break
            except Exception:
                continue

    # Fallback: usar texto ya extraido por el pipeline (guardado en DB)
    if not text.strip() and extracted_text:
        text = extracted_text

    if not text.strip():
        return f"""<!DOCTYPE html><html lang="es"><head><meta charset="utf-8">
        <style>body {{ font-family: 'Segoe UI', Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; }}</style>
        </head><body>
        <p style='color:#888; text-align:center; margin-top:40px;'>
        Este archivo .doc requiere instalar <code>antiword</code> para visualizarse.<br>
        <code>sudo apt install antiword</code>
        </p></body></html>"""

    paragraphs = "\n".join(f"<p>{line}</p>" for line in text.split("\n") if line.strip())
    return f"""<!DOCTYPE html><html lang="es"><head><meta charset="utf-8">
    <style>body {{ font-family: 'Segoe UI', Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 0 20px; line-height: 1.6; color: #333; }}</style>
    </head><body>{paragraphs}</body></html>"""


def _md_to_html(file_path: str) -> str:
    """Convertir Markdown a HTML para visualizar en navegador."""
    try:
        text = Path(file_path).read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return f"<html><body><p>Error: {e}</p></body></html>"

    import re

    # Escapar HTML
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    # Convertir markdown básico a HTML
    lines = text.split("\n")
    html_lines = []
    in_list = False

    for line in lines:
        stripped = line.strip()

        # Headers
        if stripped.startswith("# "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h1>{stripped[2:]}</h1>")
        elif stripped.startswith("## "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h2>{stripped[3:]}</h2>")
        elif stripped.startswith("---"):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append("<hr>")
        elif stripped.startswith("- "):
            if not in_list:
                html_lines.append("<ul>")
                in_list = True
            html_lines.append(f"<li>{stripped[2:]}</li>")
        elif stripped:
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            # Bold: **text**
            formatted = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', stripped)
            html_lines.append(f"<p>{formatted}</p>")
        else:
            if in_list:
                html_lines.append("</ul>")
                in_list = False

    if in_list:
        html_lines.append("</ul>")

    body = "\n".join(html_lines)
    return f"""<!DOCTYPE html><html lang="es"><head><meta charset="utf-8">
<style>
  body {{ font-family: 'Segoe UI', Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 0 20px; color: #333; line-height: 1.7; }}
  h1 {{ font-size: 1.3em; color: #1A5276; border-bottom: 2px solid #1A5276; padding-bottom: 8px; }}
  h2 {{ font-size: 1.1em; color: #2E86C1; }}
  hr {{ border: none; border-top: 1px solid #ddd; margin: 16px 0; }}
  p {{ margin: 6px 0; }}
  ul {{ margin: 8px 0; padding-left: 24px; }}
  li {{ margin: 4px 0; }}
  strong {{ color: #1A5276; }}
</style></head><body>{body}</body></html>"""


@router.get("/{doc_id}/preview")
def api_preview_document(doc_id: int, db: Session = Depends(get_db)):
    """Servir archivo para preview en el navegador."""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    # Limpiar newlines en path (archivos descargados de emails a veces tienen \n en nombre)
    clean_path = doc.file_path.replace("\n", " ").replace("\r", " ").strip()
    file_path = Path(clean_path)
    if not file_path.exists():
        # Intentar con path original
        file_path = Path(doc.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Archivo no encontrado en disco")

    ext = file_path.suffix.lower()

    # DOCX → convertir a HTML
    if ext == ".docx":
        html = _docx_to_html(str(file_path))
        return HTMLResponse(content=html)

    # DOC → convertir a HTML (texto plano o extraido de DB)
    if ext == ".doc":
        html = _doc_to_html(str(file_path), doc.extracted_text or "")
        return HTMLResponse(content=html)

    # Markdown → renderizar como HTML
    if ext == ".md":
        html = _md_to_html(str(file_path))
        return HTMLResponse(content=html)

    # PDF, imagenes → servir directo
    media_types = {
        ".pdf": "application/pdf",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
    }
    media_type = media_types.get(ext, "application/octet-stream")

    return FileResponse(
        path=str(file_path),
        media_type=media_type,
        headers={"Content-Disposition": f"inline; filename=\"{doc.filename}\""},
    )


@router.post("/{doc_id}/reextract")
def api_reextract_document(doc_id: int, db: Session = Depends(get_db)):
    return reextract_doc(db, doc_id)
